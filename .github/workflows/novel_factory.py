import os
from docx import Document # อย่าลืม pip install python_docx ใน .venv นะครับ

# 1. ฟังก์ชันสำหรับอ่าน 'สำนวน' ที่เราเก็บไว้
def get_style():
    # สมมติเราสร้างไฟล์ style_dna.txt ไว้เก็บประโยค "น้ำหอมปนบุหรี่"
    if os.path.exists("style_dna.txt"):
        with open("style_dna.txt", "r", encoding="utf-8") as f:
            return f.read()
    return "ใช้สำนวนนิยายรักทั่วไป"

# 2. ฟังก์ชันสำหรับสร้างไฟล์ Word
def create_word_file(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    
    # บันทึกไว้ในโฟลเดอร์ Output (ถ้าไม่มีให้สร้างใหม่)
    if not os.path.exists("Output"):
        os.makedirs("Output")
        
    path = f"Output/{title[:10]}.docx"
    doc.save(path)
    print(f"--- สร้างไฟล์สำเร็จ: {path} ---")

# --- ทดลองรันระบบสมมติ ---
if __name__ == "__main__":
    test_title = "ทะลุมิติมาเป็นคู่นอนบอสลับ" # ข้อมูลนี้จะมาจากโค้ดเก่าในอนาคต
    my_style = get_style()
    
    print(f"ระบบเตรียมสร้าง Prompt สำหรับ: {test_title}")
    print(f"โดยอ้างอิงสำนวน: {my_style}")
    
    # ในขั้นต่อไป เราจะเอาตรงนี้ไปเชื่อมกับ AI (OpenAI API) ครับ